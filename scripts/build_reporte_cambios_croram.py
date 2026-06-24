from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "entregables"
OUT_DIR.mkdir(exist_ok=True)
OUT_DOCX = OUT_DIR / "reporte_cambios_croram.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(96, 96, 96)
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(0, 0, 0)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_run(paragraph, text, bold=False, italic=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run


def add_para(doc, text="", style=None, bold=False, color=None, size=None, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if text:
        add_run(p, text, bold=bold, color=color, size=size)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    add_run(p, text)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        run.font.color.rgb = BLUE if level < 3 else DARK_BLUE
    return p


def add_label_detail_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_width(table, [2700, 6660])
    set_cell_margins(table)
    for idx, (label, value) in enumerate(rows):
        label_cell, value_cell = table.rows[idx].cells
        set_cell_shading(label_cell, LIGHT_GRAY)
        lp = label_cell.paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        add_run(lp, label, bold=True, color=DARK_BLUE)
        vp = value_cell.paragraphs[0]
        vp.paragraph_format.space_after = Pt(0)
        add_run(vp, value)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_changes_table(doc):
    rows = [
        ("Menu", "Se retiró Subfamilias del submenú Familias.", "Reduce opciones duplicadas para el usuario final."),
        ("Salidas", "Se sustituyó el filtro de Subfamilia por Area y se hizo obligatoria.", "Permite registrar el gasto contra un centro de consumo."),
        ("Salidas", "Se hizo obligatoria la observación de salida.", "Asegura trazabilidad mínima del movimiento."),
        ("Ordenes", "Se agregaron tokens de idempotencia y bloqueo de botones.", "Evita duplicar ordenes por doble clic o reenvío accidental."),
        ("Backend", "Se bloquearon ordenes con FOR UPDATE antes de aplicar inventario.", "Una orden recibida o confirmada no puede volver a afectar existencias."),
        ("Entradas", "El catálogo de productos muestra proveedor, fecha y precio de última compra.", "Facilita capturar nuevas compras con referencia histórica."),
        ("Salidas", "El detalle de salida muestra costo PEPS y total en pesos.", "Deja visible el costo real asociado a la salida."),
        ("Reportes", "Se agregaron detalle de entradas, detalle de salidas y consumos por area.", "Cubre reportes solicitados de unidades, precios, observaciones y costo PEPS."),
        ("Inventario", "Se reorganizó el reporte con saldo inicial, entradas, salidas y existencia.", "Elimina la columna Movimientos y mejora lectura contable."),
        ("Compras sugeridas", "Se ajustó el cálculo con consumo mensual / 30.4 * dias requeridos + tiempo de surtido.", "El reporte vuelve a generar resultados cuando hay consumo configurado o histórico."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_width(table, [1700, 4660, 3000])
    set_cell_margins(table)
    headers = ["Modulo", "Cambio implementado", "Impacto"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_run(p, header, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_run(p, value)
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(header, "Croram | Reporte técnico de cambios", color=GRAY, size=9)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "Sistema de Almacén Croram", color=GRAY, size=9)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(24)
    title.paragraph_format.space_after = Pt(4)
    add_run(title, "Reporte técnico de cambios", bold=True, color=BLACK, size=24)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    add_run(subtitle, "Sistema de Almacén Croram", color=GRAY, size=14)

    add_label_detail_table(doc, [
        ("Cliente", "Croram"),
        ("Fecha", date.today().strftime("%d/%m/%Y")),
        ("Preparado por", "Programador del sistema"),
        ("Alcance", "Cambios funcionales, validaciones, reportes y controles anti-duplicidad solicitados por el cliente."),
    ])

    add_heading(doc, "Resumen ejecutivo", 1)
    add_para(
        doc,
        "Se realizó una intervención enfocada en mejorar la operación diaria del inventario: se reforzó el registro de salidas por area, se redujo el riesgo de ordenes duplicadas, se agregaron reportes operativos y se ajustó el cálculo de compras sugeridas. Los cambios se hicieron respetando la estructura actual del proyecto PHP procedural, sin introducir migraciones de base de datos ni dependencias externas nuevas.",
    )

    add_heading(doc, "Cambios implementados", 1)
    add_changes_table(doc)

    add_heading(doc, "Detalle técnico", 1)
    add_heading(doc, "Control de duplicidad y consistencia", 2)
    add_bullet(doc, "Se agregaron tokens de solicitud en formularios de ordenes para identificar reenvíos de la misma operación.")
    add_bullet(doc, "Se deshabilitan botones de acción durante el procesamiento para evitar doble clic.")
    add_bullet(doc, "El backend valida el token y mantiene un historial corto en sesión para rechazar solicitudes repetidas.")
    add_bullet(doc, "Las acciones críticas de recibir entrada y aprobar salida bloquean la orden con SELECT ... FOR UPDATE antes de modificar inventario.")

    add_heading(doc, "Salidas e inventario PEPS", 2)
    add_bullet(doc, "Las salidas requieren area y observaciones antes de registrarse.")
    add_bullet(doc, "El costo PEPS calculado se conserva en el detalle de salida y se muestra en pantalla y reportes.")
    add_bullet(doc, "El total en pesos de una salida se calcula desde cantidad por costo PEPS.")

    add_heading(doc, "Reportes", 2)
    add_bullet(doc, "El reporte de inventario ahora presenta saldo inicial, entradas, salidas y existencia final.")
    add_bullet(doc, "Se agregaron reportes de detalle para entradas y salidas con unidades, precios, descripcion y totales.")
    add_bullet(doc, "Se agregó reporte de consumos por area y periodo con costo PEPS promedio y total.")
    add_bullet(doc, "La exportación a Excel se actualizó para incluir las nuevas columnas.")

    add_heading(doc, "Compras sugeridas", 2)
    add_bullet(doc, "El cálculo usa consumo mensual promedio dividido entre 30.4 y lo multiplica por dias de stock requeridos más tiempo de surtido.")
    add_bullet(doc, "Cuando no hay consumo histórico de 12 meses, el cálculo puede usar el consumo diario configurado en el artículo.")
    add_bullet(doc, "La pantalla mantiene un campo editable para confirmar, aumentar o disminuir el pedido sugerido.")

    add_heading(doc, "Cambios no aplicados o condicionados", 1)
    add_label_detail_table(doc, [
        ("Limpieza total de codificación", "Se corrigieron textos visibles en pantallas intervenidas. No se aplicó una conversión global agresiva porque el proyecto incluye muchas plantillas y librerías de terceros; hacerlo sin revisión visual podría introducir errores nuevos."),
        ("Saldo inicial real", "No existe un campo histórico dedicado para saldo inicial. Se calculó como existencia actual menos entradas más salidas, usando los datos disponibles."),
        ("Kardex completo", "El sistema ya cuenta con historial por artículo. Convertirlo en Kardex contable completo requiere definir formato final, reglas de saldo acumulado y si se debe persistir costo por movimiento."),
        ("Migración de seguridad", "No se movieron credenciales ni contraseñas a un esquema nuevo porque no estaba dentro del alcance operativo inmediato. Se recomienda atenderlo como fase técnica independiente."),
    ])

    add_heading(doc, "Validación realizada", 1)
    add_para(doc, "Se ejecutó validación de sintaxis PHP con PHP 8.3 de Laragon sobre los archivos propios del sistema.")
    add_label_detail_table(doc, [
        ("Resultado", "57 archivos PHP revisados sin errores de sintaxis."),
        ("Comando equivalente", "php -l sobre archivos propios, excluyendo librerías vendor, FPDF y phpqrcode."),
        ("Riesgo residual", "No se ejecutaron pruebas integradas contra base de datos productiva desde este entorno; la validación funcional final debe hacerse con usuarios y datos reales de Croram."),
    ])

    add_heading(doc, "Recomendaciones de siguiente fase", 1)
    add_bullet(doc, "Crear una configuración de base de datos fuera del código fuente.")
    add_bullet(doc, "Migrar contraseñas de texto plano a password_hash/password_verify.")
    add_bullet(doc, "Convertir consultas sensibles a prepared statements de forma progresiva.")
    add_bullet(doc, "Definir formalmente el formato de Kardex para construir saldos acumulados y costos por movimiento.")
    add_bullet(doc, "Probar el flujo completo con una orden de entrada, una salida y una exportación de reportes usando datos reales.")

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
